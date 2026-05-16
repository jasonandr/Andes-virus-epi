import docx
from docx.shared import Inches

def main():
    doc = docx.Document('drafts/andes_virus_research_letter_v13.docx')
    
    # Paragraph 40 has the image for Figure 1
    p = doc.paragraphs[40]
    p.clear()
    run = p.add_run()
    run.add_picture('figures/figure_1_combined_2x2.png', width=Inches(6.5))
    
    doc.save('drafts/andes_virus_research_letter_v14.docx')
    print("Saved v14.docx with updated Figure 1")

if __name__ == '__main__':
    main()
