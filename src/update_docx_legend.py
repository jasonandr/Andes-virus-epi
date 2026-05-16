import docx

def replace_in_paragraphs(paragraphs):
    for p in paragraphs:
        if '(C) Empirical offspring distribution aggregating' in p.text:
            p.text = p.text.replace(
                '(C) Empirical offspring distribution aggregating',
                '(C) Empirical offspring distribution plotted as side-by-side (dodged) histograms comparing'
            )

def main():
    doc = docx.Document('drafts/andes_virus_research_letter_v12.docx')
    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)
    doc.save('drafts/andes_virus_research_letter_v13.docx')
    print("Saved v13.docx")

if __name__ == '__main__':
    main()
